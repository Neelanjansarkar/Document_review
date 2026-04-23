from pathlib import Path

from langchain_core.documents import Document


class DocumentLoader:
    supported_extensions = {".pdf", ".txt", ".md", ".docx"}

    def load(self, path: Path) -> list[Document]:
        suffix = path.suffix.lower()
        if suffix not in self.supported_extensions:
            raise ValueError(f"Unsupported file type '{suffix}'. Use PDF, DOCX, TXT, or Markdown.")
        if suffix == ".pdf":
            return self._load_pdf(path)
        if suffix == ".docx":
            return self._load_docx(path)
        return self._load_text(path)

    def _load_pdf(self, path: Path) -> list[Document]:
        try:
            from pypdf import PdfReader
        except ImportError as exc:  # pragma: no cover
            raise RuntimeError("Install pypdf to ingest PDF files.") from exc

        reader = PdfReader(str(path))
        documents: list[Document] = []
        for page_number, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                documents.append(
                    Document(
                        page_content=text,
                        metadata={"source": path.name, "page": page_number, "path": str(path)},
                    )
                )
        return documents

    def _load_docx(self, path: Path) -> list[Document]:
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:  # pragma: no cover
            raise RuntimeError("Install python-docx to ingest DOCX files.") from exc

        doc = DocxDocument(str(path))
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
        return [Document(page_content=text, metadata={"source": path.name, "path": str(path)})]

    def _load_text(self, path: Path) -> list[Document]:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return [Document(page_content=text, metadata={"source": path.name, "path": str(path)})]
